#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Jan  1 20:07:28 2017

@author: sebastian
"""
from docx import Document
from docx.shared import Inches


def generar_docx(nombre_archivo, dir_imagen, NumeroId="1", NombreProducto="Reporte acueducto", Descripcion="Acueducto", Cantidad="3 ductos"):

    f = open(nombre_archivo, 'w')

    document = Document()
    document.add_heading("REPORTES ACUEDUCTO")

    document.add_heading('Contenido', level=1)
    document.add_paragraph('LISTA: ')
    p = document.add_paragraph('NUMERO ID: ')
    p.add_run(str(NumeroId)).bold = True

    p = document.add_paragraph('NOMBRE DEL PRODUCTO: ')
    p.add_run(str(NombreProducto)).bold = True

    p = document.add_paragraph('DESCRIPCIÓN DEL PRODUCTO: ')
    p.add_run(str(Descripcion)).bold = True

    p = document.add_paragraph('CANTIDAD DEL PRODUCTO: ')
    p.add_run(str(Cantidad)).bold = True

    p = document.add_paragraph('REPORTE: ')
    p.add_run("Mapa de la capa:").bold = True

    document.add_picture(dir_imagen, width=Inches(5))
    #document.add_picture('capp.png', width=Inches(1.25))
    document.save(nombre_archivo)
    return nombre_archivo
