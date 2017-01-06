#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Jan  1 20:07:28 2017

@author: sebastian
"""
from docx import Document
from docx.shared import Inches


def generar_docx(
    nombre_archivo,
    dir_imagen,
    NombreProducto="Reporte acueducto",
    numero_features="1",
    Descripcion="Acueducto",
    Cantidad="3 ductos",
    destino_documento="documentos"
    ):

    dir_archivo = destino_documento + "/" + nombre_archivo
    f = open(dir_archivo, 'w')

    document = Document()
    document.add_heading("REPORTES ACUEDUCTO")

    document.add_heading('Contenido', level=1)
    document.add_paragraph('LISTA: ')
    p = document.add_paragraph('NUMERO FEATURES: ')
    p.add_run(str(numero_features)).bold = True

    p = document.add_paragraph('NOMBRE DEL PRODUCTO: ')
    p.add_run(str(NombreProducto)).bold = True

    p = document.add_paragraph('DESCRIPCIÓN DEL PRODUCTO: ')
    p.add_run(str(Descripcion)).bold = True

    p = document.add_paragraph('CANTIDAD DEL PRODUCTO: ')
    p.add_run(str(Cantidad)).bold = True

    p = document.add_paragraph('REPORTE: ')
    p.add_run("Mapa de la capa:").bold = True

    document.add_picture(dir_imagen, width=Inches(5))
    #document.add_picture('capp.png', width=Inches(1.25))
    document.save(dir_archivo)
    return dir_archivo
